"""Inspection report export (PDF + DOCX) with the original label photo as evidence."""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

FIELD_LABELS = {
    "manufacturer": "Manufacturer Info",
    "net_quantity": "Net Quantity",
    "mrp": "MRP",
    "country_of_origin": "Country of Origin",
    "manufacture_date": "Manufacture Date",
    "consumer_care": "Consumer Care Details",
}

FIELD_ORDER = list(FIELD_LABELS.keys())

def _score(summary: Dict[str, Any]) -> int:
    total = max(int(summary.get("total_fields_checked") or 1), 1)
    return round(int(summary.get("compliant_count") or 0) / total * 100)


def _evidence_jpeg(image_bytes: Optional[bytes], max_side: int = 1200) -> Optional[bytes]:
    if not image_bytes:
        return None
    img = PILImage.open(io.BytesIO(image_bytes))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    elif img.mode == "L":
        img = img.convert("RGB")
    w, h = img.size
    longest = max(w, h)
    if longest > max_side:
        scale = max_side / longest
        img = img.resize((max(1, int(w * scale)), max(1, int(h * scale))), PILImage.Resampling.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue()


def build_pdf(
    scan: Dict[str, Any],
    image_bytes: Optional[bytes] = None,
    generated_at: Optional[datetime] = None,
) -> bytes:
    generated_at = generated_at or datetime.now(timezone.utc)
    summary = scan.get("compliance_summary") or {}
    results = scan.get("compliance_results") or {}
    structured = scan.get("structured_fields") or {}
    font = scan.get("font_analysis") or {}

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title="METRA Compliance Inspection Report",
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("MetraTitle", parent=styles["Heading1"], fontSize=16, spaceAfter=4)
    sub = ParagraphStyle("MetraSub", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#44474e"))
    body = ParagraphStyle("MetraBody", parent=styles["Normal"], fontSize=8, leading=11)
    small = ParagraphStyle("MetraSmall", parent=styles["Normal"], fontSize=7, leading=9, textColor=colors.HexColor("#44474e"))

    story: List[Any] = [
        Paragraph("METRA — Legal Metrology Inspection Report", title),
        Paragraph(
            f"Packaged Commodities Rules, 2011 · Generated {generated_at.strftime('%Y-%m-%d %H:%M UTC')}",
            sub,
        ),
        Spacer(1, 6),
    ]

    overall = summary.get("overall_status", "—")
    score = _score(summary)
    meta = [
        ["Overall status", str(overall)],
        ["Compliance score", f"{score} / 100"],
        ["Compliant / review / violations", f"{summary.get('compliant_count', 0)} / {summary.get('review_count', 0)} / {summary.get('violations_count', 0)}"],
        ["Font-size violations / reviews", f"{font.get('violations_count', 0)} / {font.get('review_count', 0)}"],
        ["DPI used", f"{font.get('dpi', '—')} ({font.get('dpi_source', '—')})"],
    ]
    meta_table = Table(meta, colWidths=[70 * mm, 100 * mm])
    meta_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f3f4f5")),
                ("BOX", (0, 0), (-1, -1), 0.3, colors.HexColor("#c4c6cf")),
                ("INNERGRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#c4c6cf")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(meta_table)
    story.append(Spacer(1, 8))

    evidence = _evidence_jpeg(image_bytes)
    if evidence:
        story.append(Paragraph("Label evidence (original photo)", styles["Heading3"]))
        img_flow = Image(io.BytesIO(evidence))
        img_flow._restrictSize(170 * mm, 90 * mm)
        story.append(img_flow)
        story.append(Spacer(1, 8))

    story.append(Paragraph("Declaration checklist", styles["Heading3"]))
    rows = [["Field", "Status", "AI / Officer Value", "Rule", "Findings"]]
    for field in FIELD_ORDER:
        r = results.get(field) or {}
        val = (structured.get(field) or {}).get("value") or "—"
        override = r.get("officer_override") or (structured.get(field) or {}).get("officer_override")
        if override and override.get("value"):
            val_cell = Paragraph(
                f"<b>AI:</b> {val}<br/><b>Officer Override:</b> {override.get('value')} "
                f"<font color='#004d40'><b>[Authoritative]</b></font>",
                small,
            )
        else:
            val_cell = Paragraph(str(val), body)
        rows.append(
            [
                Paragraph(FIELD_LABELS[field], body),
                Paragraph(str(r.get("status") or "—"), body),
                val_cell,
                Paragraph(str(r.get("rule_reference") or "—"), body),
                Paragraph(str(r.get("findings") or "—"), small),
            ]
        )
    table = Table(rows, colWidths=[32 * mm, 28 * mm, 36 * mm, 22 * mm, 52 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0b2545")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#c4c6cf")),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Font size & readability (Rule 7) — separate from missing declarations", styles["Heading3"]))
    story.append(
        Paragraph(
            str(font.get("rule_description") or "Rule 7 Table I numeral/letter minima."),
            small,
        )
    )
    font_rows = [["Field", "Status", "Height (mm)", "Required (mm)", "Findings"]]
    for field in FIELD_ORDER:
        fr = (font.get("fields") or {}).get(field) or {}
        hmm = fr.get("height_mm")
        req = fr.get("min_required_mm")
        font_rows.append(
            [
                Paragraph(FIELD_LABELS[field], body),
                Paragraph(str(fr.get("status") or "—"), body),
                Paragraph(f"{hmm:.2f}" if isinstance(hmm, (int, float)) else "—", body),
                Paragraph(f"{req:.2f}" if isinstance(req, (int, float)) else "—", body),
                Paragraph(str(fr.get("findings") or "—"), small),
            ]
        )
    font_table = Table(font_rows, colWidths=[32 * mm, 28 * mm, 26 * mm, 26 * mm, 58 * mm])
    font_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0b2545")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#c4c6cf")),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(font_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Statutory Legal Metrology references", styles["Heading3"]))
    story.append(
        Paragraph(
            "Legal Metrology Act, 2009 (Section 36(1) penalty: ₹25,000 first offence; ₹50,000 second offence; "
            "up to ₹1,00,000 and/or imprisonment up to one year for subsequent offences). "
            "Packaged Commodities Rules, 2011 (Rules 6, 7, 8, 9, 11, 12, 24).",
            small,
        )
    )

    doc.build(story)
    return buf.getvalue()


def build_docx(
    scan: Dict[str, Any],
    image_bytes: Optional[bytes] = None,
    generated_at: Optional[datetime] = None,
) -> bytes:
    generated_at = generated_at or datetime.now(timezone.utc)
    summary = scan.get("compliance_summary") or {}
    results = scan.get("compliance_results") or {}
    structured = scan.get("structured_fields") or {}
    font = scan.get("font_analysis") or {}

    doc = Document()
    title = doc.add_heading("METRA — Legal Metrology Inspection Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(
        f"Packaged Commodities Rules, 2011 · Generated {generated_at.strftime('%Y-%m-%d %H:%M UTC')}"
    )
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x44, 0x47, 0x4E)

    doc.add_paragraph(f"Overall status: {summary.get('overall_status', '—')}")
    doc.add_paragraph(f"Compliance score: {_score(summary)} / 100")
    doc.add_paragraph(
        f"Compliant / review / violations: "
        f"{summary.get('compliant_count', 0)} / {summary.get('review_count', 0)} / {summary.get('violations_count', 0)}"
    )
    doc.add_paragraph(
        f"Font-size violations / reviews: {font.get('violations_count', 0)} / {font.get('review_count', 0)} "
        f"(DPI {font.get('dpi', '—')} · {font.get('dpi_source', '—')})"
    )

    evidence = _evidence_jpeg(image_bytes)
    if evidence:
        doc.add_heading("Label evidence (original photo)", level=2)
        doc.add_picture(io.BytesIO(evidence), width=Inches(5.8))

    doc.add_heading("Declaration checklist", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Field", "Status", "AI / Officer Value", "Rule", "Findings"]):
        hdr[i].text = label
    for field in FIELD_ORDER:
        r = results.get(field) or {}
        val = (structured.get(field) or {}).get("value") or "—"
        override = r.get("officer_override") or (structured.get(field) or {}).get("officer_override")
        row = table.add_row().cells
        row[0].text = FIELD_LABELS[field]
        row[1].text = str(r.get("status") or "—")
        if override and override.get("value"):
            row[2].text = f"AI: {val}\nOfficer Override: {override.get('value')} [Authoritative]"
        else:
            row[2].text = str(val)
        row[3].text = str(r.get("rule_reference") or "—")
        row[4].text = str(r.get("findings") or "—")

    doc.add_heading("Font size & readability (Rule 7)", level=2)
    doc.add_paragraph(str(font.get("rule_description") or ""))
    ft = doc.add_table(rows=1, cols=5)
    ft.style = "Table Grid"
    for i, label in enumerate(["Field", "Status", "Height (mm)", "Required (mm)", "Findings"]):
        ft.rows[0].cells[i].text = label
    for field in FIELD_ORDER:
        fr = (font.get("fields") or {}).get(field) or {}
        hmm = fr.get("height_mm")
        req = fr.get("min_required_mm")
        row = ft.add_row().cells
        row[0].text = FIELD_LABELS[field]
        row[1].text = str(fr.get("status") or "NOT_MEASURED")
        row[2].text = "—" if hmm is None else f"{hmm:.2f}"
        row[3].text = "—" if req is None else f"{req:.1f}"
        row[4].text = str(fr.get("findings") or "—")

    doc.add_paragraph(
        "This report is decision-support for Legal Metrology officers. "
        "Physical verification on the package remains required before enforcement."
    )

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
